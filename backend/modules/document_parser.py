import os
import re
from fastapi import HTTPException


class DocumentParser:
    def extract_text(self, file_path: str, filename: str) -> str:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            text = self._extract_pdf(file_path)
        elif ext == ".docx":
            text = self._extract_docx(file_path)
        else:
            raise HTTPException(status_code=400, detail="Поддерживаются только PDF и DOCX")
        return self._clean(text)

    def _extract_pdf(self, path: str) -> str:
        try:
            import PyPDF2
            parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            return "\n".join(parts)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Ошибка чтения PDF: {e}")

    def _extract_docx(self, path: str) -> str:
        try:
            from docx import Document
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Ошибка чтения DOCX: {e}")

    def _clean(self, text: str) -> str:
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
